"""
Resume Analyzer - Flask Web Application
Extracts and categorizes technical skills from uploaded PDF/DOCX resumes.
"""

import os
import re
from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename

# PDF text extraction
import pdfplumber

# DOCX text extraction
from docx import Document

# ─────────────────────────────────────────────
# App Configuration
# ─────────────────────────────────────────────

app = Flask(__name__)

# Maximum upload size: 10 MB
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'uploads'

# Only allow PDF and DOCX files
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Create the uploads folder on startup if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)


# ─────────────────────────────────────────────
# Skills Knowledge Base
# Maps each category to a list of known keywords.
# All entries are lowercase; matching is case-insensitive.
# ─────────────────────────────────────────────

SKILLS_DB = {
    "Programming Languages": [
        "python", "java", "javascript", "typescript", "c\\+\\+", "c#", "ruby",
        "go", "golang", "rust", "swift", "kotlin", "php", "r", "matlab",
        "scala", "perl", "bash", "shell", "powershell", "dart", "lua",
        "haskell", "groovy", "julia", "elixir", "clojure", "fortran", "cobol",
        "assembly", "objective-c", "vba", "solidity",
    ],
    "Web Frameworks": [
        "react", "angular", "vue", "django", "flask", "fastapi", "spring",
        "express", "node\\.js", "nodejs", "laravel", "rails", "ruby on rails",
        "asp\\.net", "next\\.js", "nuxt\\.js", "svelte", "jquery", "bootstrap",
        "tailwind", "tailwindcss", "ember", "backbone\\.js", "spring boot",
        "struts", "hibernate", "nest\\.js", "nestjs", "hapi", "koa",
    ],
    "Databases": [
        "mysql", "postgresql", "postgres", "mongodb", "redis", "sqlite",
        "oracle", "sql server", "mssql", "firebase", "cassandra",
        "elasticsearch", "dynamodb", "mariadb", "neo4j", "couchdb",
        "influxdb", "hbase", "supabase", "cockroachdb", "snowflake",
        "bigquery", "redshift",
    ],
    "Cloud & DevOps": [
        "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
        "jenkins", "terraform", "ansible", "puppet", "chef", "github actions",
        "gitlab ci", "circleci", "travis ci", "helm", "prometheus", "grafana",
        "nginx", "apache", "linux", "unix", "vagrant", "packer", "argocd",
        "datadog", "splunk", "new relic",
    ],
    "Tools & Technologies": [
        "git", "github", "gitlab", "bitbucket", "jira", "confluence",
        "postman", "swagger", "graphql", "rest", "grpc", "kafka",
        "rabbitmq", "celery", "webpack", "babel", "eslint", "pytest",
        "junit", "selenium", "cypress", "jest", "mocha",
        "jupyter", "numpy", "pandas", "matplotlib", "scikit-learn",
        "tensorflow", "pytorch", "keras", "opencv", "hadoop", "spark",
        "tableau", "power bi", "figma", "sketch", "photoshop",
        "unity", "unreal", "blender", "jira", "trello", "asana",
    ],
}

# Human-readable display names for skills with special capitalisation
DISPLAY_NAMES = {
    "python": "Python", "java": "Java", "javascript": "JavaScript",
    "typescript": "TypeScript", "c\\+\\+": "C++", "c#": "C#",
    "ruby": "Ruby", "go": "Go", "golang": "Go (Golang)", "rust": "Rust",
    "swift": "Swift", "kotlin": "Kotlin", "php": "PHP", "r": "R",
    "matlab": "MATLAB", "scala": "Scala", "perl": "Perl", "bash": "Bash",
    "shell": "Shell", "powershell": "PowerShell", "dart": "Dart",
    "lua": "Lua", "haskell": "Haskell", "groovy": "Groovy",
    "julia": "Julia", "elixir": "Elixir", "clojure": "Clojure",
    "fortran": "Fortran", "cobol": "COBOL", "assembly": "Assembly",
    "objective-c": "Objective-C", "vba": "VBA", "solidity": "Solidity",
    "react": "React", "angular": "Angular", "vue": "Vue.js",
    "django": "Django", "flask": "Flask", "fastapi": "FastAPI",
    "spring": "Spring", "express": "Express.js", "node\\.js": "Node.js",
    "nodejs": "Node.js", "laravel": "Laravel", "rails": "Ruby on Rails",
    "ruby on rails": "Ruby on Rails", "asp\\.net": "ASP.NET",
    "next\\.js": "Next.js", "nuxt\\.js": "Nuxt.js", "svelte": "Svelte",
    "jquery": "jQuery", "bootstrap": "Bootstrap", "tailwind": "Tailwind CSS",
    "tailwindcss": "Tailwind CSS", "ember": "Ember.js",
    "backbone\\.js": "Backbone.js", "spring boot": "Spring Boot",
    "hibernate": "Hibernate", "nest\\.js": "NestJS", "nestjs": "NestJS",
    "hapi": "Hapi.js", "koa": "Koa.js",
    "mysql": "MySQL", "postgresql": "PostgreSQL", "postgres": "PostgreSQL",
    "mongodb": "MongoDB", "redis": "Redis", "sqlite": "SQLite",
    "oracle": "Oracle DB", "sql server": "SQL Server", "mssql": "SQL Server",
    "firebase": "Firebase", "cassandra": "Apache Cassandra",
    "elasticsearch": "Elasticsearch", "dynamodb": "DynamoDB",
    "mariadb": "MariaDB", "neo4j": "Neo4j", "couchdb": "CouchDB",
    "influxdb": "InfluxDB", "hbase": "HBase", "supabase": "Supabase",
    "cockroachdb": "CockroachDB", "snowflake": "Snowflake",
    "bigquery": "BigQuery", "redshift": "Redshift",
    "aws": "AWS", "azure": "Azure", "gcp": "Google Cloud (GCP)",
    "google cloud": "Google Cloud", "docker": "Docker",
    "kubernetes": "Kubernetes", "k8s": "Kubernetes (K8s)",
    "jenkins": "Jenkins", "terraform": "Terraform", "ansible": "Ansible",
    "puppet": "Puppet", "chef": "Chef", "github actions": "GitHub Actions",
    "gitlab ci": "GitLab CI", "circleci": "CircleCI",
    "travis ci": "Travis CI", "helm": "Helm", "prometheus": "Prometheus",
    "grafana": "Grafana", "nginx": "NGINX", "apache": "Apache",
    "linux": "Linux", "unix": "Unix", "vagrant": "Vagrant",
    "packer": "Packer", "argocd": "ArgoCD", "datadog": "Datadog",
    "splunk": "Splunk", "new relic": "New Relic",
    "git": "Git", "github": "GitHub", "gitlab": "GitLab",
    "bitbucket": "Bitbucket", "jira": "JIRA", "confluence": "Confluence",
    "postman": "Postman", "swagger": "Swagger", "graphql": "GraphQL",
    "rest": "REST API", "grpc": "gRPC", "kafka": "Apache Kafka",
    "rabbitmq": "RabbitMQ", "celery": "Celery", "webpack": "Webpack",
    "babel": "Babel", "eslint": "ESLint", "pytest": "PyTest",
    "junit": "JUnit", "selenium": "Selenium", "cypress": "Cypress",
    "jest": "Jest", "mocha": "Mocha", "jupyter": "Jupyter",
    "numpy": "NumPy", "pandas": "Pandas", "matplotlib": "Matplotlib",
    "scikit-learn": "Scikit-Learn", "tensorflow": "TensorFlow",
    "pytorch": "PyTorch", "keras": "Keras", "opencv": "OpenCV",
    "hadoop": "Apache Hadoop", "spark": "Apache Spark",
    "tableau": "Tableau", "power bi": "Power BI", "figma": "Figma",
    "sketch": "Sketch", "photoshop": "Photoshop", "unity": "Unity",
    "unreal": "Unreal Engine", "blender": "Blender", "trello": "Trello",
    "asana": "Asana",
}


# ─────────────────────────────────────────────
# Helper Functions
# ─────────────────────────────────────────────

def allowed_file(filename: str) -> bool:
    """Return True if the filename has an allowed extension (pdf or docx)."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(filepath: str) -> str:
    """
    Use pdfplumber to read every page of a PDF and return its full text.
    Raises ValueError if the file cannot be read.
    """
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as exc:
        raise ValueError(f"Could not read PDF file: {exc}") from exc
    return text


def extract_text_from_docx(filepath: str) -> str:
    """
    Use python-docx to read all paragraphs and table cells from a DOCX file.
    Raises ValueError if the file cannot be read.
    """
    text = ""
    try:
        doc = Document(filepath)

        # Main body paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Tables (skills are often in table format in resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as exc:
        raise ValueError(f"Could not read DOCX file: {exc}") from exc
    return text


def extract_skills(text: str) -> dict:
    """
    Scan resume text for known technical skills.

    Strategy:
      1. Lowercase the text for case-insensitive matching.
      2. For each skill keyword, use a regex word-boundary search so that,
         for example, "r" doesn't accidentally match inside "react".
      3. Map the raw keyword to its human-readable display name.
      4. Return a dict of {category: [skill, ...]} with only non-empty categories.
    """
    text_lower = text.lower()
    found_skills: dict = {}

    for category, keywords in SKILLS_DB.items():
        matched = set()
        for keyword in keywords:
            # Word-boundary pattern; keyword may already contain regex escapes
            pattern = r'(?<![a-zA-Z0-9])' + keyword + r'(?![a-zA-Z0-9])'
            if re.search(pattern, text_lower):
                display = DISPLAY_NAMES.get(keyword, keyword.replace("\\", "").title())
                matched.add(display)

        if matched:
            found_skills[category] = sorted(matched)

    return found_skills


# ─────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────

@app.route('/')
def index():
    """Serve the main upload page."""
    return render_template('index.html')


@app.route('/analyze', methods=['POST'])
def analyze():
    """
    POST /analyze — accepts a resume file, extracts text, detects skills.

    Returns JSON:
      {
        "success": true,
        "filename": "resume.pdf",
        "skills": { "Programming Languages": ["Python", ...], ... },
        "summary": { "total_skills": 12, "categories_found": 4 }
      }

    On error returns:
      { "error": "Human-readable message" }
    """

    # 1. Validate that a file was actually sent
    if 'resume' not in request.files:
        return jsonify({'error': 'No file received. Please attach a resume file.'}), 400

    file = request.files['resume']

    if file.filename == '':
        return jsonify({'error': 'No file selected. Please choose a PDF or DOCX resume.'}), 400

    if not allowed_file(file.filename):
        return jsonify({
            'error': 'Unsupported file type. Please upload a PDF (.pdf) or Word document (.docx).'
        }), 400

    filepath = None
    try:
        # 2. Save to disk temporarily (secure_filename strips directory traversal chars)
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # 3. Extract plain text depending on file type
        ext = filename.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            text = extract_text_from_pdf(filepath)
        else:
            text = extract_text_from_docx(filepath)

        # 4. Guard against empty / image-only PDFs
        if not text.strip():
            return jsonify({
                'error': (
                    'No readable text found. '
                    'The file may be a scanned image or completely empty. '
                    'Please upload a text-based resume.'
                )
            }), 400

        # 5. Detect skills
        skills = extract_skills(text)
        total = sum(len(v) for v in skills.values())

        return jsonify({
            'success': True,
            'filename': file.filename,
            'skills': skills,
            'summary': {
                'total_skills': total,
                'categories_found': len(skills),
            },
        })

    except ValueError as exc:
        return jsonify({'error': str(exc)}), 400

    except Exception as exc:
        return jsonify({'error': f'Unexpected error while processing file: {exc}'}), 500

    finally:
        # Always remove the uploaded file from disk
        if filepath and os.path.exists(filepath):
            os.remove(filepath)


# ─────────────────────────────────────────────
# Entry Point
# ─────────────────────────────────────────────

if __name__ == '__main__':
    # debug=True gives auto-reload and detailed error pages during development.
    # Never use debug=True in production.
    app.run(debug=True)
